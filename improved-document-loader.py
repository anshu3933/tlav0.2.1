# core/document_processing/document_loader.py

import os
import json
import csv
import io
import pandas as pd
from typing import List, Optional, Dict, Any, Tuple
from langchain.schema import Document
from config.logging_config import get_module_logger
from core.document_processing.document_validator import DocumentValidator

# Import specialized document loaders
from pdfminer.high_level import extract_text
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document as DocxDocument

# Create a logger for this module
logger = get_module_logger("document_loader")

class LoaderResult:
    """Stores the result of a document loading operation."""
    
    def __init__(self, 
                 success: bool,
                 document: Optional[Document] = None,
                 error_message: Optional[str] = None,
                 warning: Optional[str] = None):
        self.success = success
        self.document = document
        self.error_message = error_message
        self.warning = warning
    
    @property
    def has_warning(self) -> bool:
        """Check if result has a warning."""
        return self.warning is not None


class PDFLoader:
    """Handles loading and processing PDF documents with enhanced text extraction."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and extract text from PDF with improved extraction.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            A tuple of (extracted_text, error_message)
        """
        try:
            # First try with pdfminer's high-level interface
            text = extract_text(file_path)
            
            # Check if we extracted any text
            if text and text.strip():
                return text.strip(), None
                
            # If not, try with more detailed control over extraction parameters
            text = self._extract_with_custom_params(file_path)
            
            if not text.strip():
                return None, "Could not extract text from PDF. The file may be scanned or contain only images."
                
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return None, f"Error loading PDF: {str(e)}"
    
    def _extract_with_custom_params(self, file_path: str) -> str:
        """Extract text from PDF with custom parameters for better results.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        resource_manager = PDFResourceManager()
        output_string = io.StringIO()
        codec = 'utf-8'
        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            all_texts=True
        )
        
        with open(file_path, 'rb') as file:
            with TextConverter(resource_manager, output_string, codec=codec, laparams=laparams) as converter:
                interpreter = PDFPageInterpreter(resource_manager, converter)
                for page in PDFPage.get_pages(file, check_extractable=False):
                    interpreter.process_page(page)
                
                text = output_string.getvalue()
        
        return text


class DocxLoader:
    """Handles loading and processing DOCX documents with robust error handling."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and extract text from DOCX with error handling.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A tuple of (extracted_text, error_message)
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs with formatting preservation
            text = ""
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
                    
                    # Check for hyperlinks and other elements
                    for run in para.runs:
                        if run.hyperlink:
                            text += f" [Link: {run.hyperlink.url}]"
            
            # Extract tables
            for table in doc.tables:
                text += "\nTable Content:\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                    if row_text:
                        text += row_text + "\n"
            
            # Check if we extracted any text
            if not text.strip():
                return None, "Could not extract text from DOCX. The file may be empty."
                
            return text.strip(), None
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return None, f"Error loading DOCX: {str(e)}"


class TextLoader:
    """Handles loading text files with robust error handling."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load text file content with error handling.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            A tuple of (file_content, error_message)
        """
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'ascii', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    if content.strip():
                        return content, None
                except UnicodeDecodeError:
                    continue
            
            # If we get here, none of the encodings worked
            return None, "Could not decode text file with supported encodings."
                
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            return None, f"Error loading text file: {str(e)}"


class CSVLoader:
    """Handles loading and processing CSV files."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and parse CSV file.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            A tuple of (parsed_content, error_message)
        """
        try:
            # Try to detect dialect and encoding
            with open(file_path, 'rb') as f:
                sample = f.read(4096)
            
            # Try to detect encoding
            encodings = ['utf-8', 'latin-1', 'ascii', 'cp1252']
            detected_encoding = None
            
            for encoding in encodings:
                try:
                    sample.decode(encoding)
                    detected_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not detected_encoding:
                return None, "Could not detect CSV encoding."
            
            # Parse CSV
            df = pd.read_csv(file_path, encoding=detected_encoding)
            
            # Convert to text representation
            text = "CSV Data:\n"
            
            # Add headers
            text += " | ".join(df.columns) + "\n"
            text += "-" * 50 + "\n"
            
            # Add sample rows (first 10)
            sample_rows = df.head(10).astype(str)
            for _, row in sample_rows.iterrows():
                text += " | ".join(row.values) + "\n"
            
            # Add summary if there are more rows
            if len(df) > 10:
                text += f"\n... and {len(df) - 10} more rows.\n"
            
            # Add statistics for numeric columns
            text += "\nSummary Statistics:\n"
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                stats = df[numeric_cols].describe().to_string()
                text += stats + "\n"
            
            return text, None
            
        except Exception as e:
            logger.error(f"Error loading CSV {file_path}: {str(e)}")
            return None, f"Error loading CSV: {str(e)}"


class JSONLoader:
    """Handles loading and processing JSON files."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and parse JSON file.
        
        Args:
            file_path: Path to the JSON file
            
        Returns:
            A tuple of (parsed_content, error_message)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to formatted text representation
            if isinstance(data, list):
                text = f"JSON Data (Array with {len(data)} items):\n"
                
                # Sample the first few items
                sample_size = min(5, len(data))
                for i in range(sample_size):
                    text += f"Item {i+1}:\n"
                    text += json.dumps(data[i], indent=2) + "\n\n"
                
                if len(data) > sample_size:
                    text += f"... and {len(data) - sample_size} more items.\n"
            else:
                text = "JSON Data (Object):\n"
                text += json.dumps(data, indent=2)
            
            return text, None
            
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {file_path}: {str(e)}")
            return None, f"Invalid JSON: {str(e)}"
        except Exception as e:
            logger.error(f"Error loading JSON {file_path}: {str(e)}")
            return None, f"Error loading JSON: {str(e)}"


class ExcelLoader:
    """Handles loading and processing Excel files."""
    
    def load(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Load and parse Excel file.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            A tuple of (parsed_content, error_message)
        """
        try:
            # Read all sheets
            xl = pd.ExcelFile(file_path)
            sheet_names = xl.sheet_names
            
            text = f"Excel File with {len(sheet_names)} sheets:\n\n"
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                text += f"Sheet: {sheet_name} ({len(df)} rows, {len(df.columns)} columns)\n"
                
                # Add headers
                text += " | ".join(df.columns) + "\n"
                text += "-" * 50 + "\n"
                
                # Add sample rows (first 5)
                sample_rows = df.head(5).astype(str)
                for _, row in sample_rows.iterrows():
                    text += " | ".join(row.values) + "\n"
                
                # Add summary if there are more rows
                if len(df) > 5:
                    text += f"\n... and {len(df) - 5} more rows.\n"
                
                text += "\n"
            
            return text, None
            
        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {str(e)}")
            return None, f"Error loading Excel: {str(e)}"


class DocumentLoader:
    """Main document loading coordinator with support for multiple file types."""
    
    def __init__(self):
        """Initialize with components."""
        self.validator = DocumentValidator()
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DocxLoader(),
            '.txt': TextLoader(),
            '.csv': CSVLoader(),
            '.json': JSONLoader(),
            '.xlsx': ExcelLoader(),
            '.xls': ExcelLoader(),
            '.md': TextLoader(),  # Treat markdown as text
        }
    
    def load_documents(self, file_paths: List[str]) -> List[LoaderResult]:
        """Load multiple documents with validation and detailed results.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of LoaderResult objects
        """
        results = []
        
        for file_path in file_paths:
            result = self.load_single_document(file_path)
            results.append(result)
                
        return results
    
    def load_single_document(self, file_path: str) -> LoaderResult:
        """Load a single document with validation and detailed result.
        
        Args:
            file_path: Path to the file
            
        Returns:
            LoaderResult object
        """
        # Validate file path
        is_valid, error_message = self.validator.validate_file_path(file_path)
        if not is_valid:
            logger.error(f"File validation failed for {file_path}: {error_message}")
            return LoaderResult(success=False, error_message=error_message)
        
        try:
            # Get file extension
            extension = os.path.splitext(file_path)[1].lower()
            
            # Get appropriate loader
            loader = self.loaders.get(extension)
            if not loader:
                error_message = f"No loader found for extension: {extension}"
                logger.error(error_message)
                return LoaderResult(success=False, error_message=error_message)
            
            # Load content
            content, load_error = loader.load(file_path)
            if load_error:
                return LoaderResult(success=False, error_message=load_error)
            
            # Validate content
            is_valid, content_error = self.validator.validate_content(content)
            if not is_valid:
                return LoaderResult(success=False, error_message=content_error)
            
            # Create document with warning if any
            warning = content_error if is_valid and content_error else None
            document = Document(
                page_content=content,
                metadata=self._create_metadata(file_path, extension)
            )
            
            return LoaderResult(success=True, document=document, warning=warning)
            
        except Exception as e:
            logger.error(f"Unexpected error loading document {file_path}: {str(e)}", exc_info=True)
            return LoaderResult(
                success=False, 
                error_message=f"Unexpected error loading document: {str(e)}"
            )
    
    def _create_metadata(self, file_path: str, extension: str) -> Dict[str, Any]:
        """Create metadata for document with enhanced information.
        
        Args:
            file_path: Path to the file
            extension: File extension
            
        Returns:
            Document metadata
        """
        # Get file stats
        file_size = os.path.getsize(file_path)
        modified_time = os.path.getmtime(file_path)
        created_time = os.path.getctime(file_path)
        
        # Set document type based on extension
        doc_type = "unknown"
        if extension in ['.pdf']:
            doc_type = "pdf"
        elif extension in ['.docx', '.doc']:
            doc_type = "document"
        elif extension in ['.txt', '.md']:
            doc_type = "text"
        elif extension in ['.csv']:
            doc_type = "spreadsheet"
        elif extension in ['.xlsx', '.xls']:
            doc_type = "spreadsheet"
        elif extension in ['.json']:
            doc_type = "data"
        
        # Create metadata
        return {
            "source": file_path,
            "file_type": extension,
            "document_type": doc_type,
            "file_name": os.path.basename(file_path),
            "file_size": file_size,
            "file_size_mb": round(file_size / (1024 * 1024), 2),
            "last_modified": modified_time,
            "created": created_time,
            "id": f"doc_{os.path.basename(file_path)}_{int(modified_time)}"
        }
